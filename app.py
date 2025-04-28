import pandas as pd
import qrcode
import os
from PIL import Image
from docx import Document
from docx.shared import Inches

file_path = "data.xlsx"
df = pd.read_excel(file_path)

print(df.head())

if "Name" not in df.columns or "Link" not in df.columns:
    raise ValueError("Ensure the Excel file contains 'Name' and 'Link' columns.")

image_dir = "Images"
docx_dir = "Docx"
os.makedirs(image_dir, exist_ok=True)
os.makedirs(docx_dir, exist_ok=True)

for index, row in df.iterrows():
    name = str(row["Name"]).strip()
    link = str(row["Link"]).strip()

    if not link.startswith("http"):
        print(f"Skipping invalid link for {name}: {link}")
        continue

    print(f"Generating QR for {name} -> {link}")

    qr = qrcode.QRCode(box_size=10, border=4)
    qr.add_data(link)
    qr.make(fit=True)
    img = qr.make_image(fill="black", back_color="white")

    img_filename = os.path.join(image_dir, f"{name}.png")
    img.save(img_filename)
    print(f"Saved QR: {img_filename}")

    doc = Document()
    doc.add_picture(img_filename, width=Inches(2))
    doc.add_paragraph("\n")
    doc_path = os.path.join(docx_dir, f"{name}.docx")
    doc.save(doc_path)

print("QR codes saved successfully in Images folder and Word document in Docx folder.")
